import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import os

def pdf_to_word(pdf_path, docx_path):
    # Open the PDF
    pdf_document = fitz.open(pdf_path)
    
    # Create a new Word document
    doc = Document()

    # Loop through each page
    for page_num in range(len(pdf_document)):
        page = pdf_document.load_page(page_num)
        
        # Extract text from the page
        text = page.get_text("text")
        if text.strip():
            doc.add_paragraph(text)
        
        # Extract images from the page
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image_ext = base_image["ext"]  # Get the image extension
            image_filename = f"image_{page_num+1}_{img_index+1}.{image_ext}"
            
            # Save the image
            with open(image_filename, "wb") as image_file:
                image_file.write(image_bytes)
            
            # Insert the image into the Word document
            doc.add_picture(image_filename, width=Inches(4.0))
            
            # Optionally delete the image file after adding it to the Word document
            os.remove(image_filename)
    
    # Save the Word document
    doc.save(docx_path)
    print(f"Conversion successful! Word file saved to {docx_path}")

# Example usage
pdf_path = r"C:\Users\sudha\Downloads\Ragini_Kumari_CV.pdf"
docx_path = r"C:\Users\sudha\Downloads\Ragini_Kumari_.docx"
pdf_to_word(pdf_path, docx_path)
